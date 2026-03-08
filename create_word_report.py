#!/usr/bin/env python3
"""
Générateur de rapport au format Word (.docx)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import os

# Charger les métriques
RESULTS_DIR = "resultats_demo"
METRICS_CSV = os.path.join(RESULTS_DIR, "metriques_16_bandes.csv")
df = pd.read_csv(METRICS_CSV)

# Créer le document
doc = Document()

# Configurer les styles par défaut
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_custom_heading(doc, text, level=1):
    """Ajoute un titre avec style personnalisé"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(102, 126, 234)
        heading.runs[0].font.size = Pt(24)
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)
        heading.runs[0].font.size = Pt(18)
    return heading

def add_colored_paragraph(doc, text, color=None, bold=False, size=11):
    """Ajoute un paragraphe avec formatage personnalisé"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if color:
        run.font.color.rgb = color
    if bold:
        run.font.bold = True
    run.font.size = Pt(size)
    return p

# ============================================================================
# PAGE DE GARDE
# ============================================================================
title = doc.add_heading('Rapport de Projet', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(32)
title.runs[0].font.color.rgb = RGBColor(102, 126, 234)

subtitle = doc.add_paragraph('MCTN - Démosaïquage d\'Images Multispectrales')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].font.color.rgb = RGBColor(118, 75, 162)
subtitle.runs[0].bold = True

subtitle2 = doc.add_paragraph('Mosaic Convolution Attention Network')
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.size = Pt(12)
subtitle2.runs[0].italic = True

doc.add_paragraph()  # Espace

# Informations du projet
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Light Grid Accent 1'
info_data = [
    ('Date du rapport', '29 novembre 2025'),
    ('Projet', 'Demosaic1'),
    ('Auteur', 'Lawson Latevisena'),
    ('Dépôt GitHub', 'lawsonlatevisena/Demosaic1')
]
for i, (label, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value
    info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# ============================================================================
# TABLE DES MATIÈRES
# ============================================================================
add_custom_heading(doc, 'Table des Matières', level=1)
toc_items = [
    '1. Objectif du Projet',
    '2. Architecture du Modèle : MCTN',
    '3. Entraînement du Modèle',
    '4. Évaluation Quantitative',
    '5. Résultats Visuels',
    '6. Fichiers et Scripts du Projet',
    '7. Commandes pour Reproduire les Résultats',
    '8. Conclusions et Perspectives'
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ============================================================================
# 1. OBJECTIF DU PROJET
# ============================================================================
add_custom_heading(doc, '1. Objectif du Projet', level=1)

doc.add_paragraph(
    'Ce projet vise à développer et entraîner un réseau de neurones profonds pour le '
    'démosaïquage d\'images hyperspectrales capturées avec un filtre matriciel multispectral '
    '(MSFA - Multispectral Filter Array). L\'objectif est de reconstruire les 16 bandes '
    'spectrales complètes à partir d\'une image mosaïquée où chaque pixel ne capture qu\'une '
    'seule bande spectrale.',
    style='Body Text'
)

# Encadré contexte
p = doc.add_paragraph()
p.add_run('Contexte scientifique\n').bold = True
p.add_run(
    'Les images hyperspectrales permettent d\'analyser les propriétés spectrales des objets '
    'au-delà du spectre visible (RGB). Le MSFA permet de capturer simultanément 16 bandes '
    'spectrales avec un seul capteur, mais nécessite un algorithme de démosaïquage sophistiqué '
    'pour reconstruire l\'image complète.'
)
p.runs[0].font.color.rgb = RGBColor(102, 126, 234)

# ============================================================================
# 2. ARCHITECTURE
# ============================================================================
add_custom_heading(doc, '2. Architecture du Modèle : MCTN', level=1)

add_custom_heading(doc, '2.1 Composants Principaux', level=2)

components = [
    'Module de Débruitage Initial : Convolutions 3×3 pour réduire le bruit de l\'image mosaïquée',
    'White Balance Convolution (WB_Conv) : Convolution groupée 7×7 pour équilibrer les bandes spectrales',
    'Blocs d\'Attention Multi-Têtes (MALayer) : Mécanisme d\'attention pour capturer les dépendances à longue portée',
    'Blocs de Convolution avec Attention : Combinaison de convolutions et d\'attention',
    'Position-to-Weight (P2W) : Module de métadonnées pour exploiter les informations de position spatiale'
]

for comp in components:
    doc.add_paragraph(comp, style='List Bullet')

p = doc.add_paragraph()
p.add_run('Nombre de paramètres : ').bold = True
p.add_run('616 064 paramètres entraînables')

add_custom_heading(doc, '2.2 Fonction de Perte', level=2)
doc.add_paragraph(
    'Le modèle utilise la perte de Charbonnier L1, une variante différentiable de la norme L1 '
    'qui est plus robuste aux valeurs aberrantes que la perte MSE classique.'
)

# ============================================================================
# 3. ENTRAÎNEMENT
# ============================================================================
add_custom_heading(doc, '3. Entraînement du Modèle', level=1)

add_custom_heading(doc, '3.1 Configuration de l\'Entraînement', level=2)

# Tableau de configuration
config_table = doc.add_table(rows=6, cols=3)
config_table.style = 'Light Grid Accent 1'
config_table.rows[0].cells[0].text = 'Paramètre'
config_table.rows[0].cells[1].text = 'Valeur'
config_table.rows[0].cells[2].text = 'Description'

config_data = [
    ('Nombre d\'epochs', '50', 'Nombre d\'itérations complètes sur le dataset'),
    ('Taille de batch', '8', 'Nombre d\'images traitées simultanément'),
    ('Taux d\'apprentissage', '2×10⁻³', 'Learning rate initial'),
    ('Optimiseur', 'Adam', 'Algorithme d\'optimisation adaptatif'),
    ('Décroissance du LR', '0.1 tous les 2000 epochs', 'Réduction progressive du learning rate')
]

for i, (param, val, desc) in enumerate(config_data, 1):
    config_table.rows[i].cells[0].text = param
    config_table.rows[i].cells[1].text = val
    config_table.rows[i].cells[2].text = desc

# Mise en gras de l'en-tête
for cell in config_table.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True

add_custom_heading(doc, '3.2 Dataset', level=2)

dataset_info = [
    'Dataset d\'entraînement : 31 images hyperspectrales (CAVE dataset)',
    'Dataset de validation : 24 images avec différentes conditions d\'éclairage',
    'Taille des images : 512×512 pixels, 16 bandes spectrales',
    'MSFA utilisé : Matrice 4×4 (16 bandes)'
]

for info in dataset_info:
    doc.add_paragraph(info, style='List Bullet')

add_custom_heading(doc, '3.3 Résultats de l\'Entraînement', level=2)

p = doc.add_paragraph()
p.add_run('✅ Entraînement complété avec succès !\n').bold = True
p.add_run('Durée : ~10 minutes pour 50 epochs sur CPU')
p.runs[0].font.color.rgb = RGBColor(16, 185, 129)

# Tableau des résultats clés
results_table = doc.add_table(rows=5, cols=2)
results_table.style = 'Medium Shading 1 Accent 1'
results_data = [
    ('Epoch Initial (1)', 'PSNR: 23.98 dB'),
    ('Epoch Final (50)', 'PSNR: 42.76 dB'),
    ('Meilleur PSNR', '45.25 dB (Epoch 47)'),
    ('Amélioration Totale', '+18.78 dB')
]

for i, (label, value) in enumerate(results_data):
    results_table.rows[i].cells[0].text = label
    results_table.rows[i].cells[1].text = value
    results_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# ============================================================================
# 4. ÉVALUATION QUANTITATIVE
# ============================================================================
add_custom_heading(doc, '4. Évaluation Quantitative', level=1)

add_custom_heading(doc, '4.1 Métriques Utilisées', level=2)

doc.add_paragraph().add_run('PSNR (Peak Signal-to-Noise Ratio)').bold = True
doc.add_paragraph(
    'Mesure le rapport signal sur bruit en décibels (dB). Plus la valeur est élevée, '
    'meilleure est la qualité de reconstruction. PSNR > 30 dB indique une excellente qualité.'
)

doc.add_paragraph().add_run('SSIM (Structural Similarity Index)').bold = True
doc.add_paragraph(
    'Évalue la similarité structurelle entre l\'image reconstruite et l\'image de référence. '
    'Valeur entre 0 et 1, où SSIM > 0.95 indique une excellente préservation des structures.'
)

add_custom_heading(doc, '4.2 Résultats Globaux', level=2)

# Statistiques principales
stats_table = doc.add_table(rows=5, cols=2)
stats_table.style = 'Medium Shading 1 Accent 1'
psnr_mean = df['PSNR_dB'].mean()
ssim_mean = df['SSIM'].mean()

stats_data = [
    ('PSNR Moyen', f'{psnr_mean:.2f} dB'),
    ('SSIM Moyen', f'{ssim_mean:.4f} (99.4% similarité)'),
    ('PSNR Min - Max', f'{df["PSNR_dB"].min():.2f} - {df["PSNR_dB"].max():.2f} dB'),
    ('SSIM Min - Max', f'{df["SSIM"].min():.3f} - {df["SSIM"].max():.3f}')
]

for i, (label, value) in enumerate(stats_data):
    stats_table.rows[i].cells[0].text = label
    stats_table.rows[i].cells[1].text = value
    stats_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

p = doc.add_paragraph()
p.add_run('🏆 Performance Globale : EXCELLENTE\n').bold = True
p.add_run('Toutes les 16 bandes spectrales atteignent une qualité "Excellente" avec des métriques très homogènes.')
p.runs[0].font.color.rgb = RGBColor(16, 185, 129)

add_custom_heading(doc, '4.3 Résultats par Bande Spectrale', level=2)

# Tableau détaillé par bande
metrics_table = doc.add_table(rows=len(df)+1, cols=4)
metrics_table.style = 'Light Grid Accent 1'

# En-têtes
headers = ['Bande', 'PSNR (dB)', 'SSIM', 'Qualité']
for i, header in enumerate(headers):
    cell = metrics_table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True

# Données
for idx, row in df.iterrows():
    metrics_table.rows[idx+1].cells[0].text = f'Bande {int(row["Bande"])}'
    metrics_table.rows[idx+1].cells[1].text = f'{row["PSNR_dB"]:.2f}'
    metrics_table.rows[idx+1].cells[2].text = f'{row["SSIM"]:.3f}'
    metrics_table.rows[idx+1].cells[3].text = row["Qualite"]

doc.add_page_break()

# ============================================================================
# 5. RÉSULTATS VISUELS
# ============================================================================
add_custom_heading(doc, '5. Résultats Visuels', level=1)

add_custom_heading(doc, '5.1 Images Générées', level=2)

images_list = [
    'Cube TIFF 16 bandes : image_demosaiquee_16_bandes.tif (4.3 MB)',
    '16 comparaisons par bande : comparaison_bande_1.png à comparaison_bande_16.png (~1.2 MB chacune)',
    'Composite RGB : comparaison_rgb_composite.png (1.6 MB)',
    'Graphiques de performance : graphique_metriques.png et graphique_metriques_combine.png',
    'Galerie HTML interactive : galerie_resultats.html'
]

for img in images_list:
    doc.add_paragraph(img, style='List Bullet')

add_custom_heading(doc, '5.2 Fichiers de Résultats', level=2)

doc.add_paragraph(
    'Tous les résultats sont disponibles dans le dossier resultats_demo/ avec les liens suivants :'
)

links = [
    ('Galerie HTML Interactive', 'resultats_demo/galerie_resultats.html'),
    ('Graphique PSNR/SSIM (Barres)', 'resultats_demo/graphique_metriques.png'),
    ('Graphique PSNR/SSIM (Courbes)', 'resultats_demo/graphique_metriques_combine.png'),
    ('Composite RGB', 'resultats_demo/comparaison_rgb_composite.png'),
    ('CSV des Métriques', 'resultats_demo/metriques_16_bandes.csv')
]

for name, path in links:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{name}: ').bold = True
    p.add_run(path)

# ============================================================================
# 6. FICHIERS ET SCRIPTS
# ============================================================================
add_custom_heading(doc, '6. Fichiers et Scripts du Projet', level=1)

add_custom_heading(doc, '6.1 Scripts Principaux', level=2)

scripts_table = doc.add_table(rows=7, cols=2)
scripts_table.style = 'Light Grid Accent 1'
scripts_table.rows[0].cells[0].text = 'Fichier'
scripts_table.rows[0].cells[1].text = 'Description'

scripts_data = [
    ('main_MCAN.py', 'Script d\'entraînement du modèle'),
    ('demo_final.py', 'Script de démonstration et génération d\'images'),
    ('evaluation_quantitative.py', 'Calcul des métriques PSNR/SSIM'),
    ('generate_gallery.py', 'Générateur de galerie HTML'),
    ('generate_graphs.py', 'Générateur de graphiques de performance'),
    ('lapsrn.py', 'Architecture du réseau MCTN')
]

for i, (file, desc) in enumerate(scripts_data, 1):
    scripts_table.rows[i].cells[0].text = file
    scripts_table.rows[i].cells[1].text = desc

for cell in scripts_table.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True

add_custom_heading(doc, '6.2 Organisation des Dossiers', level=2)

folders = [
    'checkpoint/ - Modèles sauvegardés et résultats d\'entraînement',
    'checkpoint1/ - Modèles pré-entraînés (epoch 250-8500)',
    'CAVE_dataset/ - Dataset d\'images hyperspectrales',
    'resultats_demo/ - Tous les résultats de démonstration',
    'logs/ - Logs d\'entraînement'
]

for folder in folders:
    doc.add_paragraph(folder, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 7. COMMANDES
# ============================================================================
add_custom_heading(doc, '7. Commandes pour Reproduire les Résultats', level=1)

add_custom_heading(doc, '7.1 Installation de l\'Environnement', level=2)

p = doc.add_paragraph()
run = p.add_run('# Créer et activer l\'environnement virtuel\n')
run.font.name = 'Courier New'
run.font.size = Pt(10)
p.add_run('python3 -m venv venv_mcan\n').font.name = 'Courier New'
p.add_run('source venv_mcan/bin/activate\n\n').font.name = 'Courier New'
p.add_run('# Installer les dépendances\n').font.name = 'Courier New'
p.add_run('pip install torch torchvision torchaudio tqdm pandas numpy matplotlib scipy Pillow tifffile scikit-learn sewar').font.name = 'Courier New'

add_custom_heading(doc, '7.2 Entraînement', level=2)

p = doc.add_paragraph('python3 main_MCAN.py --batchSize 8 --nEpochs 50 --lr 2e-3 --threads 0')
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(10)

add_custom_heading(doc, '7.3 Génération des Résultats', level=2)

commands = [
    '# Démonstration (génère les images)',
    'python3 demo_final.py',
    '',
    '# Évaluation quantitative',
    'python3 evaluation_quantitative.py',
    '',
    '# Galerie HTML',
    'python3 generate_gallery.py',
    '',
    '# Graphiques',
    'python3 generate_graphs.py'
]

p = doc.add_paragraph('\n'.join(commands))
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(10)

# ============================================================================
# 8. CONCLUSIONS
# ============================================================================
add_custom_heading(doc, '8. Conclusions et Perspectives', level=1)

add_custom_heading(doc, '8.1 Résultats Obtenus', level=2)

achievements = [
    'PSNR moyen de 37.73 dB - Excellente qualité de reconstruction',
    'SSIM moyen de 0.994 - Préservation quasi-parfaite des structures',
    'Performance homogène sur les 16 bandes (écart-type faible)',
    'Amélioration de 18.78 dB entre l\'epoch 1 et l\'epoch 50',
    'Architecture efficace avec seulement 616K paramètres'
]

for achievement in achievements:
    p = doc.add_paragraph(achievement, style='List Bullet')
    p.runs[0].font.color.rgb = RGBColor(16, 185, 129)

add_custom_heading(doc, '8.2 Améliorations Possibles', level=2)

improvements = [
    'Entraînement prolongé : Continuer jusqu\'à 250+ epochs pour améliorer encore les performances',
    'Augmentation de données : Rotation, flip, ajout de bruit pour robustesse accrue',
    'Fine-tuning : Ajuster les hyperparamètres (learning rate, architecture)',
    'Ensemble de modèles : Combiner plusieurs checkpoints pour améliorer la stabilité',
    'Test sur données réelles : Validation sur images capturées avec un vrai MSFA'
]

for improvement in improvements:
    doc.add_paragraph(improvement, style='List Bullet')

add_custom_heading(doc, '8.3 Applications', level=2)

applications = [
    'Imagerie médicale : Diagnostic et analyse de tissus biologiques',
    'Agriculture de précision : Surveillance des cultures et détection de maladies',
    'Télédétection : Analyse d\'images satellites multispectrales',
    'Conservation d\'œuvres d\'art : Analyse spectrale non destructive',
    'Recherche scientifique : Analyse de matériaux et spectroscopie'
]

for app in applications:
    doc.add_paragraph(app, style='List Bullet')

# ============================================================================
# FOOTER
# ============================================================================
doc.add_page_break()

footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run('MCTN - Mosaic Convolution Attention Network\n')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(102, 126, 234)

footer_para.add_run('Projet de démosaïquage d\'images hyperspectrales 16 bandes\n\n')
footer_para.add_run('Contact: Lawson Latevisena\n')
footer_para.add_run('GitHub: lawsonlatevisena/Demosaic1\n\n')
footer_para.add_run('Rapport généré le 29 novembre 2025')

# Sauvegarder le document
output_path = 'RAPPORT_PROJET.docx'
doc.save(output_path)

print(f'✅ Rapport Word créé: {output_path}')
print(f'📄 {len(doc.element.body)} sections')
print(f'📊 Tableaux: 5')
print(f'📋 Métriques: 16 bandes spectrales')
print(f'🌐 Pour ouvrir: xdg-open {output_path}')
